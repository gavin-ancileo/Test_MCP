import json
import boto3
import os
from io import BytesIO
from datetime import datetime
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload
from docx import Document


# === Helper: Google Drive ===
def get_drive_service():
    """Get Google Drive service using SA JSON from Secrets Manager"""
    sm = boto3.client("secretsmanager", region_name="ap-southeast-1")
    secret_id = os.environ["GOOGLE_SA_SECRET"]  # ex: aap/gdrive/SA_JSON
    resp = sm.get_secret_value(SecretId=secret_id)
    sa_info = json.loads(resp["SecretString"])

    creds = service_account.Credentials.from_service_account_info(
        sa_info,
        scopes=["https://www.googleapis.com/auth/drive.file"]
    )
    return build("drive", "v3", credentials=creds)


def create_docx(text: str, title: str = "Output") -> bytes:
    """Create DOCX file in memory"""
    doc = Document()
    doc.add_heading(title, 0)

    for para in text.split("\n\n"):
        if para.strip():
            doc.add_paragraph(para)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def upload_to_drive(file_bytes: bytes, filename: str, folder_id: str):
    """Upload DOCX to Google Drive"""
    service = get_drive_service()
    file_metadata = {"name": filename, "parents": [folder_id]}
    media = MediaIoBaseUpload(BytesIO(file_bytes),
                              mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                              resumable=True)

    file = service.files().create(
        body=file_metadata,
        media_body=media,
        fields='id,webViewLink',
        supportsAllDrives=True
    ).execute()

    return file.get("webViewLink")


# === Lambda Handler ===
def lambda_handler(event, context):
    headers = {
        "Content-Type": "application/json",
        "Access-Control-Allow-Origin": "*",
        "Access-Control-Allow-Methods": "POST,OPTIONS",
        "Access-Control-Allow-Headers": "*"
    }

    if event.get("httpMethod") == "OPTIONS":
        return {"statusCode": 200, "headers": headers, "body": ""}

    try:
        body = json.loads(event.get("body", "{}"))
    except:
        body = {}

    prompt = body.get("prompt")
    if not prompt:
        return {"statusCode": 400, "headers": headers, "body": json.dumps({"error": "Missing prompt"})}

    try:
        # === 1. Call Bedrock ===
        bedrock = boto3.client("bedrock-runtime", region_name="ap-southeast-1")
        model_id = body.get("model", "amazon.nova-micro-v1:0")
        max_tokens = body.get("max_tokens", 1000)
        temperature = body.get("temperature", 0.7)

        # Build request based on model type
        if "claude" in model_id:
            # Anthropic Claude expects 'prompt'
            request_body = {
                "prompt": f"\n\nHuman: {prompt}\n\nAssistant:",
                "max_tokens_to_sample": max_tokens,
                "temperature": temperature
            }
        elif "nova" in model_id:
            # Amazon Nova expects 'messages'
            request_body = {
                "messages": [
                    {"role": "user", "content": [{"text": prompt}]}
                ],
                "inferenceConfig": {
                    "maxTokens": max_tokens,
                    "temperature": temperature
                }
            }
        else:
            # Fallback for Titan / other models
            request_body = {
                "inputText": prompt,
                "textGenerationConfig": {
                    "maxTokenCount": max_tokens,
                    "temperature": temperature
                }
            }

        print(f"Calling Bedrock model: {model_id}")
        response = bedrock.invoke_model(
            modelId=model_id,
            body=json.dumps(request_body)
        )

        resp_body = json.loads(response["body"].read())

        if "completion" in resp_body:
            text = resp_body["completion"]
        elif "output" in resp_body and "message" in resp_body["output"]:
            # Nova style output
            text = resp_body["output"]["message"]["content"][0]["text"]
        elif "content" in resp_body:
            if isinstance(resp_body["content"], list):
                text = resp_body["content"][0].get("text", "")
            else:
                text = resp_body["content"]
        else:
            text = str(resp_body)

        result = {"success": True, "model": model_id, "response": text}

        # === 2. Save to Drive (optional) ===
        folder_id = body.get("folder_id")
        if folder_id:
            try:
                ts = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
                filename = body.get("filename", f"bedrock_output_{ts}.docx")
                docx_bytes = create_docx(text, body.get("title", "Bedrock Output"))
                file_url = upload_to_drive(docx_bytes, filename, folder_id)
                result["file_url"] = file_url
                result["filename"] = filename
            except Exception as e:
                print("Drive upload error:", e)
                result["drive_error"] = str(e)

        return {"statusCode": 200, "headers": headers, "body": json.dumps(result)}

    except Exception as e:
        import traceback
        print("Error:", e)
        print(traceback.format_exc())
        return {"statusCode": 500, "headers": headers, "body": json.dumps({"error": str(e), "type": type(e).__name__})}
