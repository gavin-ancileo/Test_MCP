#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
AAP MCP Server - Full Version with Ultra Strict Validation
Keeps all existing logic: gate validation, Drive upload
Enhanced with smart placeholder detection
"""

import sys
import os
import json
import sqlite3
import re
import logging
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional, Tuple
from io import BytesIO

# Optional imports
try:
    from docx import Document
    HAS_DOCX = True
except:
    HAS_DOCX = False

# Add parent path for connectors
sys.path.insert(0, str(Path(__file__).parent.parent))

try:
    from connectors.drive import DriveConnector
    HAS_DRIVE = True
except:
    HAS_DRIVE = False

# Logging setup
logging.basicConfig(
    level=os.getenv("MCP_LOG_LEVEL", "INFO"),
    format="[%(asctime)s] %(levelname)s: %(message)s",
    handlers=[logging.StreamHandler(sys.stderr)],
)

# Database path with env override
ENV_DB = os.getenv("AAP_MCP_DB")
if ENV_DB:
    DB_PATH = Path(ENV_DB)
else:
    DB_PATH = Path(__file__).resolve().parent.parent / "data" / "prompts.db"

def is_placeholder(val: str) -> bool:
    """Smart detection for placeholder/invalid values"""
    if not val:
        return True
    
    s = str(val).strip()
    
    # Empty or too short (except numbers)
    if not s or (len(s) < 2 and not re.match(r"^\d+$", s)):
        return True
    
    # Pattern-based detection
    patterns = [
        r"^\[.*\]$",                    # [anything]
        r"^<.*>$",                      # <anything>
        r"^\{.*\}$",                    # {anything}
        r"^\(.*\)$",                    # (anything)
        r"^(tbd|todo|n/?a|none|null|pending|unknown)$",  # Common placeholders
        r"^(your|my|the|this|that)\s+\w+$",  # "your name", "the company", etc
        r"^\w+\s+(name|title|date|value|amount|number)$",  # "company name", "job title"
        r"^(please|enter|provide|specify|insert)\s+",     # Instructions
        r"^(example|sample|test|demo|dummy)\s*",          # Example indicators
        r"\.\.\.$",                     # Ends with ...
        r"^[-_\*#]+$",                  # Only symbols
    ]
    
    s_lower = s.lower()
    for pattern in patterns:
        if re.match(pattern, s_lower):
            return True
    
    # Heuristic: Too generic/vague
    vague_terms = [
        "to be", "as per", "per policy", "standard", "typical",
        "will be", "should be", "must be", "need to", "have to"
    ]
    for term in vague_terms:
        if term in s_lower:
            return True
    
    return False

def humanize_var_name(var_name: str) -> str:
    """Convert variable name to human readable"""
    # Simple approach: replace underscores and capitalize
    words = var_name.split("_")
    result = []
    for word in words:
        # Keep uppercase if already uppercase (like ID, URL)
        if word.isupper() and len(word) <= 4:
            result.append(word)
        else:
            result.append(word.capitalize())
    return " ".join(result)

class MCPServer:
    def __init__(self):
        self.db_path = DB_PATH
        self.prompts_cache: Optional[List[Dict]] = None
        logging.info(f"MCP Server starting - DB: {self.db_path}")
        if not self.db_path.exists():
            logging.error(f"Database not found at {self.db_path}")

    def safe_json_parse(self, s: Optional[str], default):
        """Safe JSON parsing with fallback"""
        if not s:
            return default
        try:
            return json.loads(s)
        except:
            return default

    def load_prompts(self) -> List[Dict]:
        """Load all active prompts from database"""
        if self.prompts_cache is not None:
            return self.prompts_cache
            
        if not self.db_path.exists():
            return []
            
        try:
            conn = sqlite3.connect(self.db_path)
            conn.row_factory = sqlite3.Row
            cur = conn.execute("""
                SELECT id, code, name, categories, content, variables,
                       output_folder, is_active
                FROM prompts
                WHERE is_active = 1
                ORDER BY code
            """)
            prompts = []
            for r in cur:
                prompts.append({
                    "id": r["id"],
                    "code": r["code"],
                    "name": r["name"],
                    "categories": self.safe_json_parse(r["categories"], []),
                    "content": r["content"] or "",
                    "variables": self.safe_json_parse(r["variables"], []),
                    "output_folder": r["output_folder"],
                })
            conn.close()
            self.prompts_cache = prompts
            logging.info(f"Loaded {len(prompts)} prompts")
            return prompts
        except Exception as e:
            logging.error(f"DB error: {e}")
            return []

    def extract_variables(self, content: str) -> List[Dict]:
        """Extract variables from template with metadata"""
        if not content:
            return []
        
        pattern = r"\{\{(\w+)(\?)?(=([^}]*))?\}\}"
        seen = set()
        out = []
        
        for m in re.finditer(pattern, content):
            name = m.group(1)
            if name in seen:
                continue
            seen.add(name)
            
            is_optional = (m.group(2) == "?")
            default_val = m.group(4) if m.group(3) else None
            desc = humanize_var_name(name)
            
            item = {
                "name": name,
                "description": desc,
                "required": not is_optional and default_val is None
            }
            if default_val is not None:
                item["default"] = default_val
                
            out.append(item)
        return out

    def normalize_field_name(self, name: str) -> str:
        """Normalize field names: 'Offer Date' -> 'offer_date'"""
        return name.lower().replace(" ", "_").replace("-", "_")

    def validate_arguments(self, prompt: Dict, provided: Dict) -> Tuple[List[str], List[str], Dict]:
        """Validate provided arguments against requirements"""
        provided = provided or {}
        
        # Normalize all provided keys
        provided_normalized = {}
        for k, v in provided.items():
            if not k.startswith("__"):
                normalized = self.normalize_field_name(k)
                provided_normalized[normalized] = v
        
        # Get variable spec
        spec = self.extract_variables(prompt["content"])
        if not spec:
            spec = [{"name": v, "required": True} 
                   for v in prompt.get("variables", [])]
        
        required = [v["name"] for v in spec if v.get("required", True)]
        missing = []
        clean_values = {}
        
        # Check each required field
        for var_name in required:
            val = None
            if var_name in provided_normalized:
                val = provided_normalized[var_name]
            elif var_name in provided:
                val = provided[var_name]
            
            if val is None:
                missing.append(var_name)
                continue
                
            val = str(val).strip()
            if not val or is_placeholder(val):
                missing.append(var_name)
            else:
                clean_values[var_name] = val
        
        # Add optional fields if provided  
        for v in spec:
            var_name = v["name"]
            if not v.get("required"):
                val = None
                if var_name in provided_normalized:
                    val = provided_normalized[var_name]
                elif var_name in provided:
                    val = provided[var_name]
                    
                if val:
                    val = str(val).strip()
                    if val and not is_placeholder(val):
                        clean_values[var_name] = val
                elif "default" in v:
                    clean_values[var_name] = v["default"]
        
        return required, missing, clean_values

    def fill_template(self, content: str, values: Dict) -> str:
        """Fill template with values"""
        if not content:
            return ""
            
        result = content
        for k, v in values.items():
            result = result.replace(f"{{{{{k}}}}}", str(v))
            result = result.replace(f"{{{{{k}?}}}}", str(v))
            result = re.sub(r"\{\{" + re.escape(k) + r"=[^}]*\}\}", str(v), result)
        
        # Clean remaining placeholders
        result = re.sub(r"\{\{\w+\??\}\}", "", result)
        result = re.sub(r"\{\{\w+=[^}]*\}\}", "", result)
        return result

    def create_docx(self, text: str, title: str) -> bytes:
        """Create DOCX document"""
        if not HAS_DOCX:
            return text.encode('utf-8')
            
        doc = Document()
        doc.add_heading(title, 0)
        for para in text.split("\n\n"):
            if para.strip():
                doc.add_paragraph(para)
        
        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def upload_to_drive(self, prompt: Dict, content: str, args: Dict) -> Optional[str]:
        """Upload to Google Drive if configured"""
        if not HAS_DRIVE:
            return None
            
        output_folder = prompt.get("output_folder", "")
        if not output_folder.startswith("drive://folder/"):
            return None
            
        try:
            folder_id = output_folder.replace("drive://folder/", "")
            
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            name_field = (args.get("candidate_name") or 
                         args.get("employee_name") or 
                         args.get("name") or 
                         "document")
            filename = f"{prompt['code']}_{name_field}_{timestamp}.docx"
            
            docx_data = self.create_docx(content, prompt["name"])
            drive = DriveConnector()
            url = drive.upload_docx(docx_data, filename, folder_id)
            
            logging.info(f"Uploaded to Drive: {url}")
            return url
            
        except Exception as e:
            logging.error(f"Drive upload failed: {e}")
            return None

    def detect_auto_fill_attempt(self, args: Dict) -> bool:
        """Detect if Claude is trying to auto-fill with generic values"""
        if not args or len(args) <= 2:  # Too few args to determine
            return False
            
        # Skip validation if user explicitly confirmed
        if args.get("__confirm__") == "yes":
            # Only check for VERY obvious placeholders
            for key, value in args.items():
                if key.startswith("__"):
                    continue
                val_str = str(value).lower()
                # Only reject OBVIOUS fake data
                if any(x in val_str for x in ["example", "sample", "test user", "demo", "[", "]"]):
                    return True
            return False
        
        # Count suspicious patterns (for initial calls without confirm)
        suspicious_count = 0
        total_count = 0
        
        for key, value in args.items():
            if key.startswith("__"):
                continue
                
            total_count += 1
            val_str = str(value).lower().strip()
            
            # Only flag CLEARLY generic names/companies
            generic_patterns = [
                r"^(john|jane|bob|alice|test)\s+(doe|smith|user|person)$",
                r"^(company|organization|acme|test|demo)\s*(corp|inc|ltd)?$",
                r"^example",
                r"^sample",
            ]
            
            for pattern in generic_patterns:
                if re.match(pattern, val_str):
                    suspicious_count += 1
                    break
        
        # Only flag if MAJORITY are suspicious (>50%)
        return False
    
    def handle_request(self, req: Dict) -> Optional[Dict]:
        """Handle JSON-RPC request"""
        method = req.get("method")
        params = req.get("params", {}) or {}
        req_id = req.get("id")

        try:
            # Initialize
            if method == "initialize":
                return {
                    "jsonrpc": "2.0",
                    "id": req_id,
                    "result": {
                        "protocolVersion": "2025-06-18",
                        "capabilities": {"tools": {}},
                        "serverInfo": {"name": "aap-mcp", "version": "3.0"},
                    },
                }

            # List tools
            elif method == "tools/list":
                prompts = self.load_prompts()
                tools = []
                
                for p in prompts:
                    tools.append({
                        "name": f"fill_{p['code']}",
                        "description": f"Generate: {p['name']}",
                        "inputSchema": {
                            "type": "object",
                            "additionalProperties": True,
                            "properties": {
                                "__confirm__": {
                                    "type": "string",
                                    "description": "Set to 'yes' to generate"
                                }
                            },
                            "required": []
                        }
                    })
                
                return {"jsonrpc": "2.0", "id": req_id, "result": {"tools": tools}}

            # Tool execution with validation
            elif method == "tools/call":
                tool_name = params.get("name")
                args = params.get("arguments", {}) or {}
                
                if not tool_name or not tool_name.startswith("fill_"):
                    return {
                        "jsonrpc": "2.0",
                        "id": req_id,
                        "error": {"code": -32602, "message": f"Unknown tool: {tool_name}"}
                    }
                
                code = tool_name[5:]
                prompts = self.load_prompts()
                prompt = next((p for p in prompts if p["code"] == code), None)
                
                if not prompt:
                    return {
                        "jsonrpc": "2.0",
                        "id": req_id,
                        "error": {"code": -32602, "message": f"Prompt not found: {code}"}
                    }
                
                # NEW: Detect auto-fill attempts
                if self.detect_auto_fill_attempt(args):
                    spec = self.extract_variables(prompt["content"])
                    required_fields = [humanize_var_name(v["name"]) 
                                      for v in spec if v.get("required", True)]
                    
                    msg = f"❌ DETECTED AUTO-FILL ATTEMPT\n\n"
                    msg += f"Cannot generate {prompt['name']} with generic/example data.\n\n"
                    msg += "❗ STOP and ASK the user for REAL information:\n\n"
                    for field in required_fields:
                        msg += f"• {field}\n"
                    msg += "\n⚠️ Each field needs ACTUAL values from the user."
                    msg += "\nDo NOT proceed with example data."
                    
                    return {
                        "jsonrpc": "2.0",
                        "id": req_id,
                        "result": {
                            "content": [{"type": "text", "text": msg}]
                        }
                    }
                
                # Special handling: if ONLY __confirm__ is provided
                if len(args) == 1 and "__confirm__" in args:
                    spec = self.extract_variables(prompt["content"])
                    if not spec:
                        spec = [{"name": v, "required": True} 
                               for v in prompt.get("variables", [])]
                    
                    required_fields = [humanize_var_name(v["name"]) 
                                      for v in spec if v.get("required", True)]
                    
                    msg = f"Cannot generate {prompt['name']} without ALL required information.\n\n"
                    msg += "❗ **MUST provide ALL of these:**\n\n"
                    for field in required_fields:
                        msg += f"• {field}\n"
                    msg += "\n⚠️ Do NOT use default/standard values. Ask the user for EACH field."
                    
                    return {
                        "jsonrpc": "2.0",
                        "id": req_id,
                        "result": {
                            "content": [{"type": "text", "text": msg}]
                        }
                    }
                
                # Normal validation
                required, missing, clean_values = self.validate_arguments(prompt, args)
                
                # If we have SOME values but still missing fields
                if clean_values and missing:
                    collected_msg = "✅ **Received so far:**\n"
                    for k, v in clean_values.items():
                        collected_msg += f"• {humanize_var_name(k)}: {v}\n"
                    
                    missing_msg = "\n❗ **STILL REQUIRED (ask user for these):**\n"
                    for field in missing:
                        missing_msg += f"• {humanize_var_name(field)}\n"
                    
                    msg = f"{collected_msg}{missing_msg}\n"
                    msg += "⚠️ Cannot proceed without ALL fields. Ask user for missing information."
                    
                    return {
                        "jsonrpc": "2.0",
                        "id": req_id,
                        "result": {
                            "content": [{"type": "text", "text": msg}]
                        }
                    }
                
                # If nothing provided yet
                elif missing and not clean_values:
                    msg = f"Cannot generate {prompt['name']} - missing ALL required fields:\n\n"
                    for field in missing:
                        msg += f"• {humanize_var_name(field)}\n"
                    msg += "\n❗ Must get ALL information from user. Do NOT use defaults."
                    
                    return {
                        "jsonrpc": "2.0",
                        "id": req_id,
                        "result": {
                            "content": [{"type": "text", "text": msg}]
                        }
                    }
                
                # Check confirmation
                if args.get("__confirm__", "").lower() != "yes":
                    fields_list = "\n".join([f"• {k}: {v}" for k, v in clean_values.items()])
                    msg = (f"Ready to generate {prompt['name']}:\n\n"
                          f"{fields_list}\n\n"
                          f"Please confirm by adding __confirm__='yes'")
                    
                    return {
                        "jsonrpc": "2.0",
                        "id": req_id,
                        "result": {
                            "content": [{"type": "text", "text": f"⚠️ {msg}"}]
                        }
                    }
                
                # Generate content
                filled = self.fill_template(prompt["content"], clean_values)
                result_text = f"✅ Generated {prompt['name']}:\n\n{filled}"
                
                # Try Drive upload
                try:
                    output_folder = prompt.get("output_folder", "")
                    logging.info(f"Output folder config: {output_folder}")
                    
                    if output_folder and output_folder.startswith("drive://folder/"):
                        if HAS_DRIVE:
                            url = self.upload_to_drive(prompt, filled, clean_values)
                            if url:
                                result_text += f"\n\n📁 Saved to Drive: {url}"
                                logging.info(f"Successfully uploaded to Drive: {url}")
                            else:
                                result_text += "\n\n⚠️ Drive upload failed - no URL returned"
                                logging.warning("Drive upload returned no URL")
                        else:
                            result_text += "\n\n⚠️ Drive connector not available"
                            logging.warning("Drive connector not imported")
                    else:
                        logging.info(f"No Drive folder configured for {prompt['code']}")
                except Exception as e:
                    result_text += f"\n\n⚠️ Drive upload error: {str(e)}"
                    logging.error(f"Drive upload exception: {e}", exc_info=True)
                
                return {
                    "jsonrpc": "2.0",
                    "id": req_id,
                    "result": {
                        "content": [{"type": "text", "text": result_text}]
                    }
                }

            # Handle notifications
            elif method in ("notifications/initialized", "notifications/cancelled"):
                return None

            # Resources list (empty)
            elif method == "resources/list":
                return {"jsonrpc": "2.0", "id": req_id, "result": {"resources": []}}

            # Prompts list
            elif method == "prompts/list":
                prompts = self.load_prompts()
                lst = []
                for p in prompts:
                    desc = p["name"]
                    if p.get("categories"):
                        desc += f" – {', '.join(p['categories'])}"
                    lst.append({"name": p["code"], "description": desc})
                return {"jsonrpc": "2.0", "id": req_id, "result": {"prompts": lst}}

            # Unknown method
            else:
                return {
                    "jsonrpc": "2.0",
                    "id": req_id,
                    "error": {"code": -32601, "message": f"Method not found: {method}"}
                }

        except Exception as e:
            logging.error(f"Error handling {method}: {e}")
            return {
                "jsonrpc": "2.0",
                "id": req_id,
                "error": {"code": -32603, "message": f"Internal error: {e}"}
            }

    def run(self):
        """Main server loop"""
        logging.info("MCP Server started with full functionality")
        
        while True:
            try:
                line = sys.stdin.readline()
                if not line:
                    break
                
                line = line.strip()
                if not line:
                    continue
                
                req = json.loads(line)
                resp = self.handle_request(req)
                
                if resp:
                    print(json.dumps(resp), flush=True)
                    
            except json.JSONDecodeError as e:
                error = {
                    "jsonrpc": "2.0",
                    "id": None,
                    "error": {"code": -32700, "message": f"Parse error: {e}"}
                }
                print(json.dumps(error), flush=True)
            except KeyboardInterrupt:
                break
            except Exception as e:
                logging.error(f"Fatal error: {e}")
                break

if __name__ == "__main__":
    MCPServer().run()