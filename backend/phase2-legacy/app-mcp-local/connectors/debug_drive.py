#!/usr/bin/env python3
"""
Debug script to test Google Drive connection
Run from project root directory
"""
import sys
from pathlib import Path

# Ensure we can import from connectors
current_dir = Path(__file__).parent
sys.path.insert(0, str(current_dir))

def test_drive_connection():
    print("Testing Google Drive connection...")
    
    try:
        from connectors.drive import DriveConnector
        from docx import Document
        from io import BytesIO
        
        # Initialize Drive connector
        drive = DriveConnector()
        
        if not drive.service:
            print("❌ Drive service not initialized")
            
            # Check if credentials file exists
            creds_path = current_dir / "connectors" / "drive_credentials.json"
            if creds_path.exists():
                print(f"✅ Credentials file found at: {creds_path}")
            else:
                print(f"❌ Credentials file NOT found at: {creds_path}")
                return False
            
            return False
            
        print("✅ Drive service initialized successfully")
        
        # Test creating a simple DOCX
        print("Creating test DOCX...")
        doc = Document()
        doc.add_heading("Test Contract", 0)
        doc.add_paragraph("This is a test contract for Gavin")
        doc.add_paragraph("Position: Software Developer")
        doc.add_paragraph("Salary: MYR 8500")
        
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        docx_bytes = buf.read()
        print(f"✅ DOCX created, size: {len(docx_bytes)} bytes")
        
        # Test upload to Drive
        folder_id = "1ViGWBhDzjbNit-oy0YTMaB8PSke8PVrs"
        filename = "test_contract_gavin.docx"
        
        print(f"Uploading to Drive folder: {folder_id}")
        print(f"Filename: {filename}")
        
        file_url = drive.upload_docx(docx_bytes, filename, folder_id)
        
        if file_url:
            print(f"✅ Upload successful!")
            print(f"File URL: {file_url}")
            return True
        else:
            print("❌ Upload failed - no URL returned")
            return False
            
    except ImportError as e:
        print(f"❌ Import error: {e}")
        print("Make sure you're running from the project root directory")
        return False
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()
        return False

def check_database():
    print("\nChecking database...")
    
    try:
        import sqlite3
        import json
        
        db_path = current_dir / "data" / "prompts.db"
        
        if not db_path.exists():
            print(f"❌ Database not found at: {db_path}")
            return False
            
        print(f"✅ Database found at: {db_path}")
        
        conn = sqlite3.connect(db_path)
        conn.row_factory = sqlite3.Row
        
        # List all prompts
        rows = conn.execute("SELECT code, name, output_folder FROM prompts WHERE is_active=1").fetchall()
        print(f"📊 Found {len(rows)} active prompts:")
        
        for row in rows:
            print(f"   - {row['code']}: {row['name']}")
            if row['output_folder']:
                print(f"     Output: {row['output_folder']}")
        
        # Check contract_generate specifically
        row = conn.execute("SELECT * FROM prompts WHERE code='contract_generate'").fetchone()
        
        if row:
            prompt = dict(row)
            print(f"\n✅ Found contract_generate prompt:")
            print(f"   Name: {prompt['name']}")
            print(f"   Output folder: {prompt['output_folder']}")
            print(f"   Active: {prompt['is_active']}")
            
            try:
                variables = json.loads(prompt.get('variables', '[]'))
                print(f"   Variables: {variables}")
            except:
                print(f"   Variables (raw): {prompt.get('variables', 'N/A')}")
                
        else:
            print("\n❌ contract_generate prompt not found in database")
            print("Need to run: python update_contract_prompt.py")
            
        conn.close()
        return True
        
    except Exception as e:
        print(f"❌ Database error: {e}")
        return False

def check_file_structure():
    print("\nChecking file structure...")
    
    required_files = [
        "connectors/drive.py",
        "connectors/drive_credentials.json", 
        "connectors/__init__.py",
        "data/prompts.db",
        "mcp/app_mcp.py"
    ]
    
    for file_path in required_files:
        full_path = current_dir / file_path
        if full_path.exists():
            print(f"✅ {file_path}")
        else:
            print(f"❌ {file_path} - NOT FOUND")

if __name__ == "__main__":
    print("🚀 AAP Drive Debug Tool")
    print(f"📁 Running from: {current_dir}")
    print("=" * 50)
    
    # Check file structure
    check_file_structure()
    
    # Check database
    db_ok = check_database()
    
    # Test Drive connection
    drive_ok = test_drive_connection()
    
    print(f"\n📊 Results:")
    print(f"   Database: {'✅' if db_ok else '❌'}")
    print(f"   Drive: {'✅' if drive_ok else '❌'}")
    
    if not drive_ok:
        print(f"\n🔧 Next steps:")
        print(f"   1. Make sure all files exist (check above)")
        print(f"   2. Run: pip install python-docx google-api-python-client google-auth")
        print(f"   3. Verify Google Service Account has access to the Drive folder")
        print(f"   4. If database missing contract_generate, run: python update_contract_prompt.py")