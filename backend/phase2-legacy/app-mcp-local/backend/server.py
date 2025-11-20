#!/usr/bin/env python3
from fastapi import FastAPI, HTTPException
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Optional
import json
import re
from pathlib import Path
from backend.database import Database
from connectors.drive import DriveConnector
from docx import Document
from io import BytesIO
from datetime import datetime

# Create FastAPI app
app = FastAPI(title="AAP Prompt Management API")

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize database
db = Database()

# Pydantic models
class PromptCreate(BaseModel):
    code: str
    name: str
    categories: List[str]
    content: str
    output_folder: Optional[str] = ""

class PromptTest(BaseModel):
    prompt_code: str
    variables: dict
    generate_document: Optional[bool] = True

class OrchestrateRequest(BaseModel):
    input: str

# Serve admin dashboard
admin_path = Path(__file__).parent.parent / "admin"
app.mount("/static", StaticFiles(directory=admin_path), name="static")

@app.get("/")
async def serve_admin():
    return FileResponse(admin_path / "index.html")

# API Routes
@app.get("/api/prompts")
async def get_prompts():
    prompts = db.get_all()
    return {"prompts": prompts}

@app.post("/api/prompts")
async def create_prompt(prompt: PromptCreate):
    success = db.create(
        prompt.code,
        prompt.name,
        prompt.categories,
        prompt.content,
        prompt.output_folder
    )
    if success:
        return {"message": "Prompt created successfully"}
    else:
        raise HTTPException(status_code=400, detail="Failed to create prompt (code might exist)")

@app.put("/api/prompts/{code}")
async def update_prompt(code: str, prompt: PromptCreate):
    """Update an existing prompt"""
    try:
        # Delete old prompt
        db.delete(code)

        # Create new one with updated data
        success = db.create(
            prompt.code,
            prompt.name,
            prompt.categories,
            prompt.content,
            prompt.output_folder
        )

        if success:
            return {"message": "Prompt updated successfully"}
        else:
            raise HTTPException(status_code=400, detail="Failed to update prompt")

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error updating prompt: {str(e)}")

@app.delete("/api/prompts/{code}")
async def delete_prompt(code: str):
    db.delete(code)
    return {"message": "Prompt deleted successfully"}

@app.get("/api/prompts/{code}")
async def get_prompt(code: str):
    prompt = db.get_one(code)
    if prompt:
        return prompt
    else:
        raise HTTPException(status_code=404, detail="Prompt not found")

@app.post("/api/test-prompt")
async def test_prompt(request: PromptTest):
    """Test a prompt and optionally generate document"""
    try:
        # Get prompt from database
        prompt = db.get_one(request.prompt_code)
        if not prompt:
            raise HTTPException(status_code=404, detail="Prompt not found")

        # Fill variables in content
        content = prompt['content']
        for var, value in request.variables.items():
            content = content.replace(f"{{{{{var}}}}}", str(value))

        result = {
            "success": True,
            "filled_content": content,
            "prompt_name": prompt['name']
        }

        # Generate and upload document if requested
        if request.generate_document and prompt.get('output_folder'):
            try:
                output_folder = prompt['output_folder']
                if output_folder.startswith("drive://folder/"):
                    folder_id = output_folder.replace("drive://folder/", "")

                    # Create DOCX
                    doc = Document()
                    doc.add_heading(prompt['name'], 0)

                    # Add content paragraphs
                    for para in content.split("\n\n"):
                        if para.strip():
                            doc.add_paragraph(para.strip())

                    # Convert to bytes
                    buf = BytesIO()
                    doc.save(buf)
                    buf.seek(0)
                    docx_bytes = buf.read()

                    # Upload to Drive
                    drive = DriveConnector()
                    timestamp = datetime.utcnow().strftime('%Y%m%d_%H%M%S')
                    filename = f"{request.prompt_code}_{timestamp}.docx"

                    file_url = drive.upload_docx(docx_bytes, filename, folder_id)

                    if file_url:
                        result["document_uploaded"] = True
                        result["document_url"] = file_url
                        result["message"] = f"Document generated and uploaded to Drive: {file_url}"
                    else:
                        result["document_uploaded"] = False
                        result["message"] = "Content generated but failed to upload to Drive"
                else:
                    result["document_uploaded"] = False
                    result["message"] = "Content generated (no Drive folder configured)"
            except Exception as drive_error:
                result["document_uploaded"] = False
                result["drive_error"] = str(drive_error)
                result["message"] = f"Content generated but Drive upload failed: {drive_error}"

        return result

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error testing prompt: {str(e)}")

@app.post("/api/orchestrate")
async def orchestrate(request: OrchestrateRequest):
    """Analyze user input and suggest best prompt"""
    user_input = request.input.lower()
    prompts = db.get_all()

    # Simple keyword matching
    scores = []
    for prompt in prompts:
        score = 0
        keywords = prompt['name'].lower().split() + prompt['categories']

        for keyword in keywords:
            if keyword.lower() in user_input:
                score += 10

        # Check if prompt variables are mentioned
        for var in prompt.get('variables', []):
            if var.lower() in user_input:
                score += 5

        if score > 0:
            scores.append((score, prompt))

    if not scores:
        return {"message": "No matching prompts found"}

    # Get best match
    best_prompt = sorted(scores, reverse=True)[0][1]

    # Extract variables from user input
    extracted = {}
    missing = []

    for var in best_prompt.get('variables', []):
        # Simple extraction - could be improved with NLP
        pattern = rf"{var}[:\s]+([^\s,\.]+)"
        match = re.search(pattern, user_input, re.IGNORECASE)
        if match:
            extracted[var] = match.group(1)
        else:
            missing.append(var)

    return {
        "selected": best_prompt,
        "extracted": extracted,
        "missing": missing
    }

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
