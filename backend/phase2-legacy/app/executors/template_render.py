from typing import Dict, Any
import io, json
from docx import Document
from .google_drive import (
    _dl_bytes, _download_text as _dl_text,
    upload_bytes_to_drive, convert_docx_to_google_doc, export_google_doc_to_pdf_bytes
)

# --- stubs PDF: bạn có thể thay bằng lib PyPDF sau ---
def fill_pdf_acroform(raw: bytes, values: Dict[str,Any]) -> bytes: return raw
def fill_pdf_by_labels(raw: bytes, mapping: Dict[str,str], values: Dict[str,Any]) -> bytes: return raw
# -----------------------------------------------------

def _docx_replace(template_bytes: bytes, values: Dict[str, Any]) -> bytes:
    import re
    TOK = re.compile(r"\{\{\s*([a-zA-Z0-9_]+)\s*\}\}")
    def repl_text(text: str) -> str:
        def repl(m): return str(values.get(m.group(1), m.group(0)))
        return TOK.sub(repl, text)
    doc = Document(io.BytesIO(template_bytes))
    for p in doc.paragraphs:
        for r in p.runs:
            r.text = repl_text(r.text)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.text = repl_text(r.text)
    out = io.BytesIO(); doc.save(out); out.seek(0)
    return out.read()

def _render_docx_asset(asset_uri: str, args: Dict[str, Any], ofolder: str, base_name: str, output_format: str) -> Dict[str, str]:
    tpl_id = asset_uri.split("/")[-1]
    tpl_bytes = _dl_bytes(tpl_id)
    rendered = _docx_replace(tpl_bytes, args)
    docx_id = upload_bytes_to_drive(ofolder, base_name + ".docx", rendered,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    arts = {"docx": f"drive://file/{docx_id}"}
    if output_format in ("pdf", "both"):
        gdoc_id = convert_docx_to_google_doc(docx_id)
        pdf_bytes = export_google_doc_to_pdf_bytes(gdoc_id)
        pdf_id = upload_bytes_to_drive(ofolder, base_name + ".pdf", pdf_bytes, "application/pdf")
        arts["pdf"] = f"drive://file/{pdf_id}"
    return arts

import json, boto3
from typing import Dict, Any

def run_template_renderer(assets: Dict[str, Any], args: Dict[str, Any], rules: Dict[str, Any]) -> Dict[str, Any]:
    pp = (rules or {}).get("post_process", {}) if rules else {}
    ofolder = args.get("output_folder_id") or pp.get("default_output_folder_id")
    if not ofolder:
        raise RuntimeError("Missing output folder: provide args.output_folder_id or rules.post_process.default_output_folder_id")

    def render_name(tpl: str) -> str:
        out = tpl or "document"
        for k, v in (args or {}).items():
            out = out.replace("${" + k + "}", str(v))
        return out

    documents = (rules or {}).get("documents") or []
    if not documents:
        for k in assets.keys():
            documents = [{
                "name": k,
                "asset_key": k,
                "strategy": "auto",
                "output_format": "both",
                "naming": pp.get("file_naming", "document")
            }]
            break

    artifacts = {}
    for doc in documents:
        name = doc.get("name", "document")
        asset_key = doc["asset_key"]
        asset = assets.get(asset_key) or {}
        mime = (asset.get("mime_type") or "").lower()
        strategy = doc.get("strategy") or (rules.get("renderer") or "auto")
        base_name = render_name(doc.get("naming") or pp.get("file_naming") or name)

        if strategy == "auto":
            if "wordprocessingml.document" in mime or asset_key.endswith("_docx"):
                strategy = "docx_tokens"
            elif "pdf" in mime:
                strategy = "pdf_acroform"
            else:
                raise RuntimeError(f"Cannot auto-resolve strategy for mime={mime}")

        if strategy == "docx_tokens":
            arts = _render_docx_asset(asset["uri"], args, ofolder, base_name, doc.get("output_format", "both"))
        elif strategy == "pdf_acroform":
            arts = _render_pdf_acroform_asset(asset["uri"], args, ofolder, base_name)
        elif strategy == "pdf_labels":
            labels_key = doc.get("labels_asset_key") or "pdf_labels_json"
            labels_asset = assets.get(labels_key)
            if not labels_asset:
                raise RuntimeError("labels_asset_key not present in assets for pdf_labels")
            arts = _render_pdf_labels_asset(asset["uri"], labels_asset["uri"], args, ofolder, base_name)
        elif strategy == "llm":
            # 👇 NEW: Gọi Claude v2 trên Bedrock
            prompt_text = asset.get("content", "")
            for k, v in (args or {}).items():
                prompt_text = prompt_text.replace(f"{{{{{k}}}}}", str(v))

            bedrock = boto3.client("bedrock-runtime", region_name="ap-southeast-1")
            body = {
                "prompt": f"\n\nHuman: {prompt_text}\n\nAssistant:",
                "max_tokens_to_sample": 300
            }
            resp = bedrock.invoke_model(
                modelId="anthropic.claude-v2",
                body=json.dumps(body)
            )
            payload = json.loads(resp["body"].read())
            arts = {"txt": payload.get("completion") or str(payload)}
        else:
            raise RuntimeError(f"Unknown strategy: {strategy}")

        for ext, uri in arts.items():
            artifacts[f"{name}_{ext}"] = uri

    return {"artifacts": artifacts}

