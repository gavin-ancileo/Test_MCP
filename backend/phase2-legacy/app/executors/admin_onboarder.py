import io, re
from typing import Dict, Any, List, Optional
from docx import Document
from ..db import tx, create_prompt, create_prompt_version, add_asset
from .google_drive import _drive, _dl_bytes

TOK = re.compile(r"\{\{\s*([a-zA-Z0-9_]+)\s*\}\}")

def _extract_tokens(file_id: str) -> List[str]:
    data = _dl_bytes(file_id); doc = Document(io.BytesIO(data))
    found = []
    for p in doc.paragraphs: found += TOK.findall(p.text)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                found += TOK.findall(cell.text)
    return sorted(set(found))

def _assets_from_docx_ids(file_ids: List[str]) -> Dict[str,Any]:
    svc = _drive(); assets = {}
    for fid in file_ids:
        meta = svc.files().get(fileId=fid, fields="id,name,mimeType").execute()
        key = meta["name"].lower().replace(".docx","").replace(" ","_") + "_docx"
        assets[key] = {"provider":"drive","uri":f"drive://file/{fid}","mime_type": meta["mimeType"]}
    return assets

def _build_arguments(tokens: List[str], required_vars: List[str]) -> List[Dict[str,Any]]:
    req = set(required_vars or [])
    args = [{"name":t, "description":f"Auto-detected '{t}'", "required": t in req, "type":"string"} for t in tokens]
    if not any(a["name"]=="output_format" for a in args):
        args.append({"name":"output_format","description":"docx|pdf|both","required":False,"type":"string",
                     "enum":["docx","pdf","both"],"default":"both"})
    return args

def preview_manifest_from_docx(code: str, name: str, version: int, docx_ids: List[str],
                               required_vars: List[str], default_output_folder_id: Optional[str]) -> Dict[str,Any]:
    tokens = set(); [tokens.update(_extract_tokens(fid)) for fid in docx_ids]
    assets = _assets_from_docx_ids(docx_ids)
    docs = []
    for key in assets.keys():
        docs.append({
            "name": key.replace("_docx",""),
            "asset_key": key,
            "strategy": "docx_tokens",
            "output_format":"both",
            "naming": "${candidate_name}_" + key.replace("_docx","") + "_${employment_start_date}"
        })
    rules = {
        "renderer":"auto",
        "post_process":{"file_naming":"${candidate_name}_${employment_start_date}"},
        "documents": docs
    }
    if default_output_folder_id:
        rules["post_process"]["default_output_folder_id"] = default_output_folder_id

    return {"code":code,"name":name,"version":version,
            "rules_json":rules,"assets":assets,
            "arguments":_build_arguments(sorted(tokens), required_vars)}

def commit_manifest(m: Dict[str,Any]) -> int:
    with tx() as cur:
        pid = create_prompt(cur, m["code"], m["name"])
        pvid = create_prompt_version(cur, pid, m["version"], m["arguments"], m["rules_json"], True)
        for k,v in m["assets"].items():
            add_asset(cur, pvid, k, v["provider"], v["uri"], v.get("mime_type",""))
        return pvid

def run_admin_onboarder_docx(args: Dict[str,Any]) -> Dict[str,Any]:
    manifest = preview_manifest_from_docx(
        args["prompt_code"], args["prompt_name"], int(args.get("version",1)),
        args["docx_file_ids"], args.get("required_vars", []),
        args.get("default_output_folder_id")
    )
    if args.get("confirm"):
        pvid = commit_manifest(manifest)
        return {"mode":"committed","prompt_version_id":pvid,"code":manifest["code"],"version":manifest["version"]}
    return {"mode":"preview","manifest":manifest}
